import docx

doc = docx.Document(r"c:\Users\abhishek.johri\OneDrive - SDET TECH\Documents\Elevate QA New Design\ElevateQA_Conference_Planning_Playbook_v0.1.docx")

# Write text and tables to a text file for easy inspection
with open(r"c:\Users\abhishek.johri\OneDrive - SDET TECH\Documents\Elevate QA New Design\scratch\docx_content.txt", "w", encoding="utf-8") as f:
    f.write("=== PARAGRAPHS ===\n")
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + "\n")
            
    f.write("\n=== TABLES ===\n")
    for i, table in enumerate(doc.tables):
        f.write(f"\nTable {i+1}:\n")
        for row in table.rows:
            row_data = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            f.write(" | ".join(row_data) + "\n")

print("Done! Extracted content to scratch/docx_content.txt")
